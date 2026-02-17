"""
עורך ספרים - כלי להחלפת מילים אוטומטית לפי הוצאות ספרים
"""

import streamlit as st
import json
import re
from pathlib import Path
from datetime import datetime
from copy import deepcopy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd
from io import BytesIO

# הגדרות בסיסיות
DATA_DIR = Path(__file__).parent / "data"
PUBLISHERS_FILE = DATA_DIR / "publishers.json"

# יצירת תיקיות אם לא קיימות
DATA_DIR.mkdir(exist_ok=True)

# הגדרת העמוד
st.set_page_config(
    page_title="עורך ספרים",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS מותאם לעברית
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Heebo:wght@300;400;500;700&display=swap');
    
    * {
        font-family: 'Heebo', sans-serif !important;
    }
    
    .main .block-container {
        direction: rtl;
        text-align: right;
    }
    
    h1, h2, h3, h4, h5, h6, p, div, span, label {
        direction: rtl;
        text-align: right;
    }
    
    /* ===== טאבים ===== */
    .stTabs [data-baseweb="tab-list"] {
        direction: rtl;
        gap: 8px;
        border-bottom: 2px solid #e0e0e0;
    }
    
    .stTabs [data-baseweb="tab"] {
        direction: rtl;
        font-weight: 500;
    }
    
    /* ===== שדות קלט ===== */
    .stSelectbox > div > div {
        direction: rtl;
    }
    
    .stTextInput > div > div > input {
        direction: rtl;
        text-align: right;
        border: 1.5px solid #c0c0c0;
        border-radius: 6px;
    }
    
    .stTextInput > div > div > input:focus {
        border-color: #667eea;
        box-shadow: 0 0 0 2px rgba(102, 126, 234, 0.2);
    }
    
    .stTextArea > div > div > textarea {
        direction: rtl;
        text-align: right;
        border: 1.5px solid #c0c0c0;
        border-radius: 6px;
    }
    
    .stTextArea > div > div > textarea:focus {
        border-color: #667eea;
        box-shadow: 0 0 0 2px rgba(102, 126, 234, 0.2);
    }
    
    /* ===== טבלאות - מסגרת חיצונית בולטת ===== */
    [data-testid="stDataFrame"],
    [data-testid="stDataEditor"] {
        border: 2px solid #4a5568;
        border-radius: 8px;
        overflow: hidden;
        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.12);
    }
    
    /* כותרת הטבלה */
    [data-testid="stDataFrame"] [data-testid="glideDataEditor"],
    [data-testid="stDataEditor"] [data-testid="glideDataEditor"] {
        border-radius: 6px;
    }
    
    /* ===== עמודות - קו מפריד ===== */
    [data-testid="stHorizontalBlock"] > [data-testid="stVerticalBlockBorderWrapper"]:not(:last-child) {
        border-left: 2px solid #e2e8f0;
        padding-left: 1rem;
    }
    
    /* ===== תיבות מידע מעוצבות ===== */
    .success-box {
        background-color: #d4edda;
        border: 1.5px solid #a3d9a5;
        border-radius: 8px;
        padding: 16px;
        margin: 10px 0;
        direction: rtl;
        box-shadow: 0 1px 4px rgba(0, 0, 0, 0.06);
    }
    
    .info-box {
        background-color: #e7f3ff;
        border: 1.5px solid #90bff9;
        border-radius: 8px;
        padding: 16px;
        margin: 10px 0;
        direction: rtl;
        box-shadow: 0 1px 4px rgba(0, 0, 0, 0.06);
    }
    
    .warning-box {
        background-color: #fff3cd;
        border: 1.5px solid #f0d264;
        border-radius: 8px;
        padding: 16px;
        margin: 10px 0;
        direction: rtl;
        box-shadow: 0 1px 4px rgba(0, 0, 0, 0.06);
    }
    
    /* ===== כותרת סקציה ===== */
    .section-header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 12px 20px;
        border-radius: 8px;
        margin-bottom: 16px;
        font-weight: 500;
        box-shadow: 0 2px 6px rgba(102, 126, 234, 0.3);
    }
    
    /* ===== קווים מפרידים ===== */
    hr {
        border: none;
        border-top: 1.5px solid #e2e8f0;
        margin: 1rem 0;
    }
    
    /* ===== כפתורים ===== */
    .stButton > button {
        border-radius: 6px;
        font-weight: 500;
        transition: all 0.2s ease;
    }
    
    .stButton > button:hover {
        transform: translateY(-1px);
        box-shadow: 0 2px 8px rgba(0, 0, 0, 0.15);
    }
</style>
""", unsafe_allow_html=True)


def load_publishers() -> dict:
    """טעינת נתוני הוצאות הספרים"""
    if PUBLISHERS_FILE.exists():
        with open(PUBLISHERS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}


def save_publishers(data: dict):
    """שמירת נתוני הוצאות הספרים"""
    with open(PUBLISHERS_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def parse_dictionary_file(content: str) -> list:
    """
    פענוח קובץ מילון בפורמט:
    "מילה למציאה" "מילה להחלפה"
    """
    entries = []
    lines = content.strip().split('\n')
    pattern = r'"([^"]+)"\s+"([^"]+)"'
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        match = re.match(pattern, line)
        if match:
            entries.append({"from": match.group(1), "to": match.group(2)})
    
    return entries


def parse_dictionary_file_detailed(content: str) -> list:
    """
    פענוח קובץ מילון עם זיהוי שורות תקינות ולא תקינות.
    מחזיר רשימה של כל השורות עם סטטוס תקינות.
    """
    entries = []
    lines = content.strip().split('\n')
    pattern = r'"([^"]+)"\s+"([^"]+)"'
    
    for line_num, line in enumerate(lines, start=1):
        stripped = line.strip()
        if not stripped:
            continue
        match = re.match(pattern, stripped)
        if match:
            entries.append({
                "line": line_num,
                "from": match.group(1),
                "to": match.group(2),
                "valid": True
            })
        else:
            entries.append({
                "line": line_num,
                "from": stripped,
                "to": "",
                "valid": False
            })
    
    return entries


def find_duplicate_entry(dictionary: list, from_text: str) -> int:
    """בדיקה האם ערך קיים במילון, מחזיר מספר שורה או -1"""
    for idx, entry in enumerate(dictionary):
        if entry["from"] == from_text:
            return idx + 1
    return -1


def add_to_deletion_history(publishers: dict, publisher_name: str, entries: list):
    """הוספת ערכים להיסטוריית המחיקות"""
    if "deletion_history" not in publishers[publisher_name]:
        publishers[publisher_name]["deletion_history"] = []
    
    history = publishers[publisher_name]["deletion_history"]
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    for entry in entries:
        history.insert(0, {
            "from": entry["from"],
            "to": entry["to"],
            "deleted_at": timestamp
        })
    
    # שמירת רק 100 הערכים האחרונים
    publishers[publisher_name]["deletion_history"] = history[:100]


def process_document(doc: Document, dictionary: list) -> tuple[Document, list]:
    """עיבוד מסמך Word והחלפת מילים עם סימון עקוב אחר שינויים (Track Changes)"""
    changes = []
    author = "עורך ספרים"
    date_str = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
    rev_id = 1

    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'

    def make_run(text, rpr=None, is_del_text=False):
        """יצירת אלמנט run חדש עם טקסט ועיצוב"""
        r = OxmlElement('w:r')
        if rpr is not None:
            r.append(deepcopy(rpr))
        tag = 'w:delText' if is_del_text else 'w:t'
        t = OxmlElement(tag)
        t.set(XML_SPACE, 'preserve')
        t.text = text
        r.append(t)
        return r

    def process_paragraph(paragraph, para_idx):
        nonlocal rev_id
        p_elem = paragraph._element

        # איסוף כל ה-runs מהפסקה
        run_elements = [child for child in p_elem if child.tag == f'{{{W_NS}}}r']
        if not run_elements:
            return

        # בניית מפת מיקומים: לכל run שומרים טקסט, עיצוב ומיקום בטקסט המלא
        runs_data = []
        pos = 0
        for rel in run_elements:
            t_elements = rel.findall(f'{{{W_NS}}}t')
            run_text = ''.join((t.text or '') for t in t_elements)
            rpr = rel.find(f'{{{W_NS}}}rPr')
            runs_data.append({
                'element': rel,
                'text': run_text,
                'start': pos,
                'end': pos + len(run_text),
                'rPr': deepcopy(rpr) if rpr is not None else None
            })
            pos += len(run_text)

        full_text = ''.join(rd['text'] for rd in runs_data)
        if not full_text:
            return

        # מציאת כל ההחלפות בטקסט המקורי
        replacements = []
        for entry in dictionary:
            from_text = entry["from"]
            to_text = entry["to"]
            search_start = 0
            while True:
                found = full_text.find(from_text, search_start)
                if found == -1:
                    break
                replacements.append((found, found + len(from_text), from_text, to_text))
                search_start = found + len(from_text)

        if not replacements:
            return

        # מיון לפי מיקום וסינון חפיפות
        replacements.sort()
        filtered = []
        last_end = 0
        for r in replacements:
            if r[0] >= last_end:
                filtered.append(r)
                last_end = r[1]
        replacements = filtered

        # רישום שינויים ללוג
        for _, _, from_text, to_text in replacements:
            changes.append({
                "שורה": para_idx,
                "מקור": from_text,
                "הוחלף ל": to_text
            })

        # בניית רשימת מקטעים: keep (ללא שינוי) או replace (החלפה)
        segments = []
        cur = 0
        for start, end, from_text, to_text in replacements:
            if cur < start:
                segments.append(('keep', cur, start))
            segments.append(('replace', start, end, from_text, to_text))
            cur = end
        if cur < len(full_text):
            segments.append(('keep', cur, len(full_text)))

        def get_portions(char_start, char_end):
            """קבלת חלקי runs (עיצוב + טקסט) עבור טווח תווים"""
            portions = []
            for rd in runs_data:
                o_start = max(char_start, rd['start'])
                o_end = min(char_end, rd['end'])
                if o_start < o_end:
                    txt = rd['text'][o_start - rd['start']:o_end - rd['start']]
                    portions.append((rd['rPr'], txt))
            return portions

        # מציאת נקודת הכנסה - שומר על אלמנטים לפני ה-runs (כמו pPr)
        ref_element = None
        for child in p_elem:
            if child.tag == f'{{{W_NS}}}r':
                break
            ref_element = child

        # הסרת כל ה-runs הישנים מהפסקה
        for rd in runs_data:
            p_elem.remove(rd['element'])

        # חישוב מיקום הכנסה
        if ref_element is not None:
            insert_idx = list(p_elem).index(ref_element) + 1
        else:
            insert_idx = 0

        # בניית אלמנטים חדשים לפי המקטעים
        for segment in segments:
            if segment[0] == 'keep':
                _, seg_start, seg_end = segment
                for rpr, text in get_portions(seg_start, seg_end):
                    p_elem.insert(insert_idx, make_run(text, rpr))
                    insert_idx += 1

            elif segment[0] == 'replace':
                _, seg_start, seg_end, from_text, to_text = segment

                # אלמנט מחיקה <w:del> - הטקסט המקורי עם העיצוב המקורי
                del_el = OxmlElement('w:del')
                del_el.set(qn('w:id'), str(rev_id))
                del_el.set(qn('w:author'), author)
                del_el.set(qn('w:date'), date_str)
                rev_id += 1

                del_portions = get_portions(seg_start, seg_end)
                for rpr, text in del_portions:
                    del_el.append(make_run(text, rpr, is_del_text=True))

                p_elem.insert(insert_idx, del_el)
                insert_idx += 1

                # אלמנט הוספה <w:ins> - הטקסט החדש עם עיצוב מה-run הראשון
                ins_el = OxmlElement('w:ins')
                ins_el.set(qn('w:id'), str(rev_id))
                ins_el.set(qn('w:author'), author)
                ins_el.set(qn('w:date'), date_str)
                rev_id += 1

                first_rpr = del_portions[0][0] if del_portions else None
                ins_el.append(make_run(to_text, first_rpr))

                p_elem.insert(insert_idx, ins_el)
                insert_idx += 1

    # עיבוד כל הפסקאות בגוף המסמך
    processed = set()
    para_idx = 0

    for paragraph in doc.paragraphs:
        para_idx += 1
        elem_id = id(paragraph._element)
        if elem_id not in processed:
            processed.add(elem_id)
            process_paragraph(paragraph, para_idx)

    # עיבוד פסקאות בטבלאות
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    para_idx += 1
                    elem_id = id(paragraph._element)
                    if elem_id not in processed:
                        processed.add(elem_id)
                        process_paragraph(paragraph, para_idx)

    return doc, changes


def main():
    st.title("📚 עורך הספרים של מירה רוזנפלד")
    st.markdown("##### כלי להחלפת מילים אוטומטית לפי הוצאות ספרים")
    
    publishers = load_publishers()
    
    # אתחול session state
    if "confirm_delete" not in st.session_state:
        st.session_state.confirm_delete = False
    if "confirm_clear_dictionary" not in st.session_state:
        st.session_state.confirm_clear_dictionary = False
    if "show_history" not in st.session_state:
        st.session_state.show_history = False
    
    tab1, tab2 = st.tabs(["🔄 עיבוד מסמך", "⚙️ ניהול מילונים"])
    
    # ===== טאב עיבוד מסמך =====
    with tab1:
        col1, col2 = st.columns([2, 1])
        
        with col1:
            uploaded_file = st.file_uploader(
                "📤 העלאת קובץ Word",
                type=["docx"],
                help="העלה קובץ Word מתורגם לעיבוד"
            )
        
        with col2:
            if publishers:
                selected_publisher = st.selectbox(
                    "🏢 בחירת הוצאת ספרים",
                    options=list(publishers.keys()),
                    index=None,
                    placeholder="בחר הוצאה",
                    help="בחר את הוצאת הספרים עבורה מיועד הספר"
                )
            else:
                st.warning("אין הוצאות ספרים מוגדרות. עבור לטאב 'ניהול מילונים' להוספה.")
                selected_publisher = None
        
        if uploaded_file and selected_publisher:
            st.markdown("---")
            dictionary = publishers[selected_publisher].get("dictionary", [])
            st.markdown(f"""
            <div class="info-box">
                <strong>🏢 הוצאה נבחרת:</strong> {selected_publisher}<br>
                <strong>📖 מספר כללים במילון:</strong> {len(dictionary)}
            </div>
            """, unsafe_allow_html=True)
            
            if st.button("🚀 בצע עיבוד", type="primary", use_container_width=True):
                with st.spinner("מעבד את המסמך..."):
                    doc = Document(uploaded_file)
                    processed_doc, changes = process_document(doc, dictionary)
                    
                    if changes:
                        st.markdown(f"""
                        <div class="success-box">
                            <strong>✅ העיבוד הושלם בהצלחה!</strong><br>
                            בוצעו {len(changes)} החלפות במסמך.
                        </div>
                        """, unsafe_allow_html=True)
                        
                        st.markdown("### 📊 לוג שינויים")
                        df = pd.DataFrame(changes)
                        st.dataframe(df, use_container_width=True, hide_index=True)
                        
                        output = BytesIO()
                        processed_doc.save(output)
                        output.seek(0)
                        
                        original_name = uploaded_file.name.replace(".docx", "")
                        st.download_button(
                            label="📥 הורד קובץ מעובד",
                            data=output,
                            file_name=f"{original_name}_מעובד.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary",
                            use_container_width=True
                        )
                    else:
                        st.markdown("""
                        <div class="warning-box">
                            <strong>ℹ️ לא נמצאו התאמות</strong><br>
                            לא נמצאו מילים להחלפה במסמך לפי המילון הנבחר.
                        </div>
                        """, unsafe_allow_html=True)
    
    # ===== טאב ניהול מילונים =====
    with tab2:
        st.markdown("### ⚙️ ניהול הוצאות ספרים ומילונים")
        
        col_publishers, col_dictionary = st.columns([1, 2])
        
        # ===== עמודה שמאלית: ניהול הוצאות ספרים =====
        with col_publishers:
            st.markdown('<div class="section-header">🏢 הוצאות ספרים</div>', unsafe_allow_html=True)
            
            if publishers:
                selected_for_edit = st.selectbox(
                    "בחר הוצאה",
                    options=list(publishers.keys()),
                    index=None,
                    placeholder="בחר הוצאה",
                    key="edit_publisher_select"
                )
            else:
                selected_for_edit = None
                st.info("אין הוצאות ספרים. הוסף הוצאה חדשה למטה.")
            
            st.markdown("---")
            st.markdown("**➕ הוספת הוצאה חדשה**")
            
            new_publisher_name = st.text_input(
                "שם ההוצאה", 
                key="new_publisher", 
                placeholder="לדוגמה: הוצאת כנרת"
            )
            new_publisher_desc = st.text_input(
                "תיאור (אופציונלי)", 
                key="new_publisher_desc", 
                placeholder="תיאור קצר"
            )
            
            if st.button("הוסף הוצאה", type="primary", use_container_width=True):
                if not new_publisher_name.strip():
                    st.error("יש להזין שם הוצאה")
                elif new_publisher_name in publishers:
                    st.error("הוצאה בשם זה כבר קיימת")
                else:
                    publishers[new_publisher_name] = {
                        "description": new_publisher_desc,
                        "dictionary": [],
                        "deletion_history": []
                    }
                    save_publishers(publishers)
                    st.success(f"הוצאה '{new_publisher_name}' נוספה!")
                    st.rerun()
            
            # עריכת שם הוצאה
            if selected_for_edit:
                st.markdown("---")
                st.markdown("**✏️ עריכת שם הוצאה**")
                rename_value = st.text_input(
                    "שם חדש להוצאה",
                    value=selected_for_edit,
                    key="rename_publisher",
                    label_visibility="collapsed"
                )
                if st.button("שנה שם", use_container_width=True):
                    new_name = rename_value.strip()
                    if not new_name:
                        st.error("יש להזין שם")
                    elif new_name == selected_for_edit:
                        st.info("השם לא השתנה")
                    elif new_name in publishers:
                        st.error("הוצאה בשם זה כבר קיימת")
                    else:
                        publishers[new_name] = publishers.pop(selected_for_edit)
                        save_publishers(publishers)
                        st.success(f"השם שונה ל-'{new_name}'")
                        st.rerun()
            
            # מחיקת הוצאה
            if selected_for_edit:
                st.markdown("---")
                st.markdown("**🗑️ מחיקת הוצאה**")
                
                if not st.session_state.confirm_delete:
                    if st.button("מחק הוצאה", type="secondary", use_container_width=True):
                        st.session_state.confirm_delete = True
                        st.rerun()
                else:
                    st.error(f"⚠️ האם אתה בטוח שברצונך למחוק את '{selected_for_edit}'?")
                    st.warning("פעולה זו תמחק את ההוצאה וכל המילון שלה לצמיתות!")
                    
                    col_yes, col_no = st.columns(2)
                    with col_yes:
                        if st.button("✅ כן, מחק", type="primary", use_container_width=True):
                            del publishers[selected_for_edit]
                            save_publishers(publishers)
                            st.session_state.confirm_delete = False
                            st.success("ההוצאה נמחקה!")
                            st.rerun()
                    with col_no:
                        if st.button("❌ ביטול", type="secondary", use_container_width=True):
                            st.session_state.confirm_delete = False
                            st.rerun()
        
        # ===== עמודה ימנית: ניהול מילון =====
        with col_dictionary:
            if selected_for_edit:
                st.markdown(f'<div class="section-header">📖 מילון: {selected_for_edit}</div>', unsafe_allow_html=True)
                
                publisher_data = publishers[selected_for_edit]
                dictionary = publisher_data.get("dictionary", [])
                deletion_history = publisher_data.get("deletion_history", [])
                
                # כפתור היסטוריה
                history_col, spacer_col = st.columns([1, 2])
                with history_col:
                    if deletion_history:
                        if st.button(f"🕐 היסטוריית מחיקות ({len(deletion_history)})", use_container_width=True):
                            st.session_state.show_history = not st.session_state.show_history
                            st.rerun()
                
                # הצגת היסטוריה
                if st.session_state.show_history and deletion_history:
                    st.markdown("---")
                    st.markdown("**🕐 היסטוריית מחיקות:**")
                    
                    # יצירת DataFrame עם checkbox
                    history_df = pd.DataFrame([
                        {
                            "בחר": False,
                            "#": i + 1,
                            "מקור": entry["from"],
                            "יעד": entry["to"],
                            "נמחק ב": entry["deleted_at"]
                        }
                        for i, entry in enumerate(deletion_history)
                    ])
                    
                    edited_history = st.data_editor(
                        history_df,
                        use_container_width=True,
                        height=200,
                        hide_index=True,
                        column_config={
                            "בחר": st.column_config.CheckboxColumn("בחר", width="small"),
                            "#": st.column_config.NumberColumn("#", width="small", disabled=True),
                            "מקור": st.column_config.TextColumn("מקור", disabled=True),
                            "יעד": st.column_config.TextColumn("יעד", disabled=True),
                            "נמחק ב": st.column_config.TextColumn("נמחק ב", disabled=True),
                        },
                        key="history_editor"
                    )
                    
                    restore_col, close_col = st.columns(2)
                    with restore_col:
                        if st.button("♻️ שחזר נבחרים", type="primary", use_container_width=True):
                            selected_rows = edited_history[edited_history["בחר"] == True]
                            if not selected_rows.empty:
                                restored_count = 0
                                for _, row in selected_rows.iterrows():
                                    # בדיקה שלא קיים כבר
                                    if find_duplicate_entry(dictionary, row["מקור"]) < 0:
                                        dictionary.append({"from": row["מקור"], "to": row["יעד"]})
                                        restored_count += 1
                                    
                                    # הסרה מההיסטוריה
                                    idx = int(row["#"]) - 1
                                    if idx < len(deletion_history):
                                        deletion_history[idx] = None
                                
                                # ניקוי None מההיסטוריה
                                publishers[selected_for_edit]["deletion_history"] = [
                                    h for h in deletion_history if h is not None
                                ]
                                publishers[selected_for_edit]["dictionary"] = dictionary
                                save_publishers(publishers)
                                st.success(f"שוחזרו {restored_count} ערכים!")
                                st.session_state.show_history = False
                                st.rerun()
                            else:
                                st.warning("לא נבחרו ערכים לשחזור")
                    
                    with close_col:
                        if st.button("✖️ סגור היסטוריה", use_container_width=True):
                            st.session_state.show_history = False
                            st.rerun()
                    
                    st.markdown("---")
                
                # הצגת המילון הקיים
                st.markdown("**רשימת מילים קיימת:**")
                
                if dictionary:
                    df_dict = pd.DataFrame([
                        {"#": i + 1, "יעד": entry["to"], "מקור": entry["from"]}
                        for i, entry in enumerate(dictionary)
                    ])
                    
                    edited_df = st.data_editor(
                        df_dict,
                        use_container_width=True,
                        height=300,
                        hide_index=True,
                        column_config={
                            "#": st.column_config.NumberColumn("#", width="small", disabled=True),
                            "מקור": st.column_config.TextColumn("מקור", width="medium"),
                            "יעד": st.column_config.TextColumn("יעד", width="medium"),
                        },
                        num_rows="dynamic",
                        key="dict_editor"
                    )
                    
                    st.caption(f"סה״כ {len(dictionary)} ערכים במילון")
                    
                    # הורדת רשימת מילים לקובץ
                    dict_lines = [f'"{e["from"]}" "{e["to"]}"' for e in dictionary]
                    dict_content = "\n".join(dict_lines)
                    st.download_button(
                        "📥 הורד רשימת מילים לקובץ",
                        data=dict_content.encode("utf-8"),
                        file_name=f"{selected_for_edit}_dictionary.txt",
                        mime="text/plain",
                        use_container_width=True
                    )
                    
                    # בדיקת כפילויות
                    source_values = []
                    duplicates = []
                    for idx, row in edited_df.iterrows():
                        if pd.notna(row["מקור"]) and str(row["מקור"]).strip():
                            source_val = str(row["מקור"]).strip()
                            if source_val in source_values:
                                duplicates.append(source_val)
                            source_values.append(source_val)
                    
                    if duplicates:
                        for dup_val in duplicates:
                            st.error(f"⚠️ כפילות: הערך '{dup_val}' מופיע יותר מפעם אחת ברשימה")
                    
                    save_disabled = len(duplicates) > 0
                    
                    if st.button("💾 שמור שינויים בטבלה", type="primary", use_container_width=True, disabled=save_disabled):
                        new_dictionary = []
                        for _, row in edited_df.iterrows():
                            if pd.notna(row["מקור"]) and pd.notna(row["יעד"]) and str(row["מקור"]).strip() and str(row["יעד"]).strip():
                                new_dictionary.append({
                                    "from": str(row["מקור"]).strip(),
                                    "to": str(row["יעד"]).strip()
                                })
                        
                        # בדיקת ערכים שנמחקו
                        new_sources = {e["from"] for e in new_dictionary}
                        deleted_entries = [e for e in dictionary if e["from"] not in new_sources]
                        
                        if deleted_entries:
                            add_to_deletion_history(publishers, selected_for_edit, deleted_entries)
                        
                        publishers[selected_for_edit]["dictionary"] = new_dictionary
                        save_publishers(publishers)
                        st.success("השינויים נשמרו!")
                        st.rerun()
                    
                    # מחיקת כל המילון
                    st.markdown("---")
                    if not st.session_state.confirm_clear_dictionary:
                        if st.button("🗑️ מחק את כל המילון", type="secondary", use_container_width=True):
                            st.session_state.confirm_clear_dictionary = True
                            st.rerun()
                    else:
                        st.error(f"⚠️ האם אתה בטוח שברצונך למחוק את כל המילון?")
                        st.warning(f"פעולה זו תמחק {len(dictionary)} ערכים!")
                        
                        col_yes, col_no = st.columns(2)
                        with col_yes:
                            if st.button("✅ כן, מחק הכל", type="primary", use_container_width=True, key="confirm_clear"):
                                # שמירה בהיסטוריה
                                add_to_deletion_history(publishers, selected_for_edit, dictionary)
                                publishers[selected_for_edit]["dictionary"] = []
                                save_publishers(publishers)
                                st.session_state.confirm_clear_dictionary = False
                                st.success("המילון נמחק!")
                                st.rerun()
                        with col_no:
                            if st.button("❌ ביטול", type="secondary", use_container_width=True, key="cancel_clear"):
                                st.session_state.confirm_clear_dictionary = False
                                st.rerun()
                else:
                    st.info("המילון ריק. הוסף ערכים באמצעות הטפסים למטה.")
                
                st.markdown("---")
                
                # הוספה ידנית
                st.markdown("**➕ הוספה ידנית**")
                add_col1, add_col2 = st.columns(2)
                with add_col1:
                    new_from = st.text_input("מקור (מה למצוא)", key="new_from", placeholder="הטקסט המקורי")
                with add_col2:
                    new_to = st.text_input("יעד (מה להחליף)", key="new_to", placeholder="הטקסט החדש")
                
                if st.button("הוסף למילון", key="add_to_dict", use_container_width=True):
                    if not (new_from.strip() and new_to.strip()):
                        st.error("יש למלא את שני השדות: מקור ויעד")
                    else:
                        existing_row = find_duplicate_entry(dictionary, new_from.strip())
                        
                        if existing_row > 0:
                            st.error(f"⚠️ הערך '{new_from}' כבר קיים במילון בשורה {existing_row}")
                        else:
                            dictionary.append({"from": new_from.strip(), "to": new_to.strip()})
                            publishers[selected_for_edit]["dictionary"] = dictionary
                            save_publishers(publishers)
                            st.success("הערך נוסף!")
                            st.rerun()
            else:
                st.markdown('<div class="section-header">📖 מילון</div>', unsafe_allow_html=True)
                if publishers:
                    st.info("👆 בחר הוצאת ספרים כדי לנהל את המילון שלה")
                else:
                    st.info("אין הוצאות ספרים. הוסף הוצאה חדשה בעמודה משמאל.")
            
            # === טעינה מקובץ - תמיד זמין ===
            st.markdown("---")
            st.markdown("**📁 טעינה מקובץ**")
            st.caption('כל שורה בפורמט: "מקור" "יעד"')
            
            uploaded_dict = st.file_uploader(
                "העלה קובץ מילון",
                type=["txt"],
                key="dict_file",
                label_visibility="collapsed"
            )
            
            if uploaded_dict:
                content = uploaded_dict.getvalue().decode("utf-8")
                file_entries = parse_dictionary_file_detailed(content)
                
                if not file_entries:
                    st.error("לא נמצאו שורות בקובץ")
                else:
                    valid_entries_list = [e for e in file_entries if e["valid"]]
                    invalid_entries_list = [e for e in file_entries if not e["valid"]]
                    
                    if invalid_entries_list:
                        invalid_lines_str = ", ".join(str(e["line"]) for e in invalid_entries_list)
                        st.warning(f"⚠️ {len(invalid_entries_list)} שורות לא תקינות (שורות: {invalid_lines_str}). ניתן לערוך ולתקן בטבלה.")
                    
                    st.success(f"✅ {len(valid_entries_list)} ערכים תקינים מתוך {len(file_entries)} שורות")
                    
                    # טבלת תצוגה מקדימה עם אפשרות עריכה
                    file_df = pd.DataFrame([
                        {
                            "#": e["line"],
                            "מקור": e["from"],
                            "יעד": e["to"],
                        }
                        for e in file_entries
                    ])
                    
                    edited_file_df = st.data_editor(
                        file_df,
                        use_container_width=True,
                        height=min(300, 60 + len(file_entries) * 35),
                        hide_index=True,
                        column_config={
                            "#": st.column_config.NumberColumn("#", width="small", disabled=True),
                            "מקור": st.column_config.TextColumn("מקור", width="medium"),
                            "יעד": st.column_config.TextColumn("יעד", width="medium"),
                        },
                        key="file_preview_editor"
                    )
                    
                    # הורדת קובץ מתוקן
                    corrected_lines = []
                    for _, row in edited_file_df.iterrows():
                        from_val = str(row["מקור"]).strip() if pd.notna(row["מקור"]) else ""
                        to_val = str(row["יעד"]).strip() if pd.notna(row["יעד"]) else ""
                        if from_val and to_val:
                            corrected_lines.append(f'"{from_val}" "{to_val}"')
                    
                    if corrected_lines:
                        corrected_content = "\n".join(corrected_lines)
                        st.download_button(
                            "💾 הורד קובץ מתוקן",
                            data=corrected_content.encode("utf-8"),
                            file_name=f"corrected_{uploaded_dict.name}",
                            mime="text/plain",
                            use_container_width=True
                        )
                    
                    # הוספת ערכים להוצאה
                    st.markdown("---")
                    
                    # חישוב ערכים תקינים מהטבלה הערוכה
                    entries_to_process = []
                    for _, row in edited_file_df.iterrows():
                        from_val = str(row["מקור"]).strip() if pd.notna(row["מקור"]) else ""
                        to_val = str(row["יעד"]).strip() if pd.notna(row["יעד"]) else ""
                        if from_val and to_val:
                            entries_to_process.append({"from": from_val, "to": to_val})
                    
                    invalid_in_table = len(edited_file_df) - len(entries_to_process)
                    
                    if selected_for_edit:
                        current_dict = publishers[selected_for_edit].get("dictionary", [])
                        
                        dup_entries = [e for e in entries_to_process if find_duplicate_entry(current_dict, e["from"]) >= 0]
                        new_unique_entries = [e for e in entries_to_process if find_duplicate_entry(current_dict, e["from"]) < 0]
                        
                        st.markdown(f"**📊 סיכום הוספה להוצאה '{selected_for_edit}':**")
                        if new_unique_entries:
                            st.markdown(f"✅ **{len(new_unique_entries)}** ערכים חדשים להוספה")
                        if dup_entries:
                            st.markdown(f"⚠️ **{len(dup_entries)}** ערכים כבר קיימים במילון (ידולגו)")
                            with st.expander("הצג ערכים כפולים"):
                                for d in dup_entries:
                                    st.text(f'"{d["from"]}" → "{d["to"]}"')
                        if invalid_in_table > 0:
                            st.markdown(f"❌ **{invalid_in_table}** שורות לא תקינות (ידולגו)")
                        
                        if new_unique_entries:
                            if st.button(
                                f"הוסף {len(new_unique_entries)} ערכים חדשים",
                                key="add_from_file",
                                type="primary",
                                use_container_width=True
                            ):
                                current_dict.extend(new_unique_entries)
                                publishers[selected_for_edit]["dictionary"] = current_dict
                                save_publishers(publishers)
                                st.success(f"נוספו {len(new_unique_entries)} ערכים בהצלחה!")
                                st.rerun()
                        elif entries_to_process and not new_unique_entries:
                            st.info("כל הערכים התקינים כבר קיימים במילון")
                    else:
                        st.warning("⚠️ יש לבחור הוצאה קודם, או ליצור הוצאה חדשה ולבחור אותה, כדי להוסיף ערכים למילון")


if __name__ == "__main__":
    main()
